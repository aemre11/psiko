"""
Streamlit rules (PsychStats) — apply on every change in this module:
1. Never write to widget-backed session state key inside on_change — pending flag; consume before widgets.
2. Every rerun-surviving widget needs stable key= initialized at startup.
3. Never call st.rerun() inside a callback — flag + natural rerun (button handlers excepted).
4. Clear only downstream state keys, not upstream keys backing visible widgets.
5. Loop-rendered widgets use index-stable keys.
6. Always use .get(key, default) for nested dicts in session state.
7. Never store class instances in session state.
8. On file upload, unconditionally overwrite KEY_RAW_DF.
9. List item action buttons use stable unique ID keys, not positional index keys.
"""

import datetime
import io
import platform
import re

import numpy as np
import pandas as pd
import pingouin as pg
import scipy
from scipy import stats
import streamlit as st
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from modules.correlation import (
    KEY_CORRELATION_FIGURE,
    KEY_CORRELATION_RESULTS,
    KEY_MODERATION_FIGURE,
    KEY_MODERATION_RESULTS,
)
from modules.data_manager import (
    KEY_COL_ROLES,
    KEY_COMPOSITE_CONFIG,
    KEY_FINAL_DF,
    KEY_SHORT_LABELS,
)
from modules.descriptives import KEY_NORMALITY_RESULTS
from modules.inferential import KEY_INFERENTIAL_RESULTS
from utils.formatters import (
    format_anova_narrative,
    format_ci,
    format_conditional_effects_narrative,
    format_correlation_narrative,
    format_moderation_narrative,
    format_normality_narrative,
    format_p,
    format_reliability_narrative,
    format_stat,
    format_ttest_narrative,
)

_FONT_NAME = "Times New Roman"
_FONT_SIZE_BODY = Pt(12)
_FONT_SIZE_TABLE = Pt(11)

# Single source of truth for the app version stamped into every exported report,
# so each document is traceable to the version that produced it.
APP_VERSION = "1.0.0"


def _display_name(var: str) -> str:
    return re.sub(r'_(Mean|Total|Sum)$', '', var, flags=re.IGNORECASE)


def _alpha_interpretation(alpha: float) -> str:
    if alpha >= 0.90:
        return "Mükemmel"
    if alpha >= 0.80:
        return "İyi"
    if alpha >= 0.70:
        return "Kabul Edilebilir"
    if alpha >= 0.60:
        return "Şüpheli"
    return "Yetersiz"


def _set_cell_border(cell, **edges) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge_name, border_spec in edges.items():
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), border_spec.get("val", "single"))
        edge.set(qn("w:sz"), str(border_spec.get("sz", 8)))
        edge.set(qn("w:color"), border_spec.get("color", "000000"))
        borders.append(edge)
    tc_pr.append(borders)


def _clear_table_borders(table) -> None:
    border_none = {"val": "nil", "sz": 0, "color": "auto"}
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(
                cell,
                top=border_none, bottom=border_none,
                left=border_none, right=border_none,
                insideH=border_none, insideV=border_none,
            )


def _apply_apa_header_borders(table) -> None:
    border_single = {"val": "single", "sz": 8, "color": "000000"}
    border_none = {"val": "nil", "sz": 0, "color": "auto"}
    header = table.rows[0]
    for cell in header.cells:
        _set_cell_border(cell, top=border_single, bottom=border_single,
                         left=border_none, right=border_none)
    if len(table.rows) > 1:
        for cell in table.rows[-1].cells:
            _set_cell_border(cell, bottom=border_single,
                             left=border_none, right=border_none)


def _style_table_cell(cell, *, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = _FONT_NAME
            run.font.size = _FONT_SIZE_TABLE
            run.bold = bold
        if not paragraph.runs:
            run = paragraph.add_run(paragraph.text)
            paragraph.text = ""
            run.font.name = _FONT_NAME
            run.font.size = _FONT_SIZE_TABLE
            run.bold = bold


def _make_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = _FONT_NAME
    normal.font.size = _FONT_SIZE_BODY
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)
    return doc


def _add_apa_table(doc: Document, headers: list[str], rows: list[list]) -> None:
    if not headers:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = str(header)
        _style_table_cell(cell, bold=True)
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = "" if value is None else str(value)
            _style_table_cell(cell, bold=False)
    _clear_table_borders(table)
    _apply_apa_header_borders(table)
    doc.add_paragraph()


def _add_title(doc: Document, title_text: str) -> None:
    p = doc.add_paragraph(title_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = doc.styles["Heading 1"]
    for run in p.runs:
        run.bold = False
        run.font.name = _FONT_NAME
        run.font.size = Pt(14)


def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Heading 2")
    for run in p.runs:
        run.font.name = _FONT_NAME


def _add_paragraph(doc: Document, text: str) -> None:
    if not text:
        return
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = _FONT_NAME
        run.font.size = _FONT_SIZE_BODY


def _dataframe_to_table_rows(df: pd.DataFrame) -> tuple[list[str], list[list]]:
    headers = [str(c) for c in df.columns]
    rows = [[row[c] for c in df.columns] for _, row in df.iterrows()]
    return headers, rows


def _frequency_table_for_export(series: pd.Series) -> pd.DataFrame:
    total = len(series)
    valid = int(series.notna().sum())
    rows: list[dict] = []
    for value, frequency in series.value_counts(dropna=True).items():
        freq = int(frequency)
        rows.append({
            "Değer": value,
            "Frekans": freq,
            "Yüzde": f"{(freq / total * 100) if total else 0.0:.1f}%",
            "Geçerli Yüzde": f"{(freq / valid * 100) if valid else 0.0:.1f}%",
            "Kümülatif Yüzde": "",
        })
    missing_freq = int(series.isna().sum())
    if missing_freq > 0:
        rows.append({
            "Değer": "Kayıp",
            "Frekans": missing_freq,
            "Yüzde": f"{(missing_freq / total * 100) if total else 0.0:.1f}%",
            "Geçerli Yüzde": "",
            "Kümülatif Yüzde": "",
        })
    cum = 0.0
    for r in rows:
        if r["Değer"] != "Kayıp" and valid:
            pct = float(r["Yüzde"].rstrip("%"))
            cum += pct
            r["Kümülatif Yüzde"] = f"{cum:.1f}%"
    return pd.DataFrame(rows)


def _descriptive_export_row(series: pd.Series, var_name: str) -> dict:
    clean = series.dropna()
    n = len(clean)
    if n == 0:
        return {"Değişken": var_name, "N": 0, "Ort": "—", "SS": "—",
                "Min": "—", "Max": "—", "Çarpıklık": "—", "Basıklık": "—", "95% GA (Ort)": "—"}
    skew = float(stats.skew(clean, bias=False))
    kurt = float(stats.kurtosis(clean, bias=False, fisher=True))
    if n < 2:
        ci_str = "—"
    else:
        lo, hi = stats.t.interval(0.95, df=n - 1,
                                   loc=float(clean.mean()),
                                   scale=float(clean.std(ddof=1) / np.sqrt(n)))
        ci_str = f"[{lo:.2f}, {hi:.2f}]"
    return {
        "Değişken": var_name, "N": n,
        "Ort": f"{clean.mean():.2f}",
        "SS": f"{clean.std(ddof=1):.2f}" if n > 1 else "—",
        "Min": f"{clean.min():.2f}", "Max": f"{clean.max():.2f}",
        "Çarpıklık": f"{skew:.3f}", "Basıklık": f"{kurt:.3f}",
        "95% GA (Ort)": ci_str,
    }


def _normality_decision(sw_p: float, skewness: float, kurtosis: float) -> str:
    if sw_p > 0.05 and abs(skewness) < 2 and abs(kurtosis) < 7:
        return "Normal ✓"
    return "Normal Değil ✗"


def _section_reliability(doc: Document, final_df: pd.DataFrame, composites: list) -> None:
    short_labels = st.session_state.get(KEY_SHORT_LABELS, {})
    reverse_labels = {v: k for k, v in short_labels.items()}
    rows: list[list] = []

    for comp in composites:
        comp_name = comp.get("name")
        short_cols = comp.get("cols", comp.get("columns", []))
        item_cols = []
        for sc in short_cols:
            if sc in final_df.columns:
                item_cols.append(sc)
            elif sc in reverse_labels and reverse_labels[sc] in final_df.columns:
                item_cols.append(reverse_labels[sc])
        if len(item_cols) < 2:
            continue
        try:
            alpha_val, _ = pg.cronbach_alpha(data=final_df[item_cols].dropna())
        except Exception:
            alpha_val = float("nan")
        n_used = int(final_df[item_cols].dropna().shape[0])
        rows.append([
            comp_name, len(item_cols),
            format_stat(float(alpha_val), 3),
            n_used,
            _alpha_interpretation(float(alpha_val)),
        ])
        narrative = format_reliability_narrative(comp_name, len(item_cols), float(alpha_val))
        if narrative:
            _add_paragraph(doc, narrative)

    if rows:
        _add_apa_table(doc, ["Ölçek", "Madde Sayısı", "α", "N (Kullanılan)", "Yorum"], rows)
    else:
        _add_paragraph(doc, "Güvenilirlik analizi için 2 veya daha fazla maddeli bileşik değişken bulunamadı.")


def _section_demographics(doc: Document, final_df: pd.DataFrame, col_roles: dict) -> None:
    demo_cols = [c for c in final_df.columns if col_roles.get(c) == "demographic"]
    if not demo_cols:
        _add_paragraph(doc, "Demografik sütun tanımlanmamıştır.")
        return
    for col in demo_cols:
        title = col if len(col) <= 60 else col[:57] + "…"
        _add_section_heading(doc, title)
        freq_df = _frequency_table_for_export(final_df[col])
        headers, trows = _dataframe_to_table_rows(freq_df)
        _add_apa_table(doc, headers, trows)


def _section_descriptives(doc: Document, final_df: pd.DataFrame, composites: list) -> None:
    rows: list[dict] = []
    for comp in composites:
        name = comp.get("name")
        if not name or name not in final_df.columns:
            continue
        rows.append(_descriptive_export_row(final_df[name], name))
    if rows:
        df = pd.DataFrame(rows)
        headers, trows = _dataframe_to_table_rows(df)
        _add_apa_table(doc, headers, trows)
    else:
        _add_paragraph(doc, "Son veri setinde bileşik değişken bulunamadı.")


def _section_normality(
    doc: Document,
    normality_results: dict,
    composites: list,
    final_df=None,
) -> None:
    # Compute normality inline if not available from session state
    if not normality_results and final_df is not None and composites:
        normality_results = {}
        for comp in composites:
            name = comp.get("name", "")
            if name and name in final_df.columns:
                clean = final_df[name].dropna()
                if len(clean) >= 3:
                    stat_val, p_val = stats.shapiro(clean)
                    normality_results[name] = {
                        "W": float(stat_val), "p": float(p_val),
                        "skewness": float(clean.skew()),
                        "kurtosis": float(clean.kurt()),
                        "normal": p_val > 0.05 and abs(clean.skew()) < 2 and abs(clean.kurt()) < 7,
                    }

    rows = []
    var_names = []
    for comp in composites:
        name = comp.get("name", "")
        if not name or name not in normality_results:
            continue
        result = normality_results[name]
        skewness = result.get("skewness", result.get("skew", 0.0))
        kurtosis = result.get("kurtosis", result.get("kurt", 0.0))
        sw_stat = result.get("W", result.get("sw_stat", 0.0))
        sw_p = result.get("p", result.get("sw_p", 1.0))
        rows.append([
            name, f"{skewness:.3f}", f"{kurtosis:.3f}",
            format_stat(sw_stat, 3), format_p(sw_p),
            _normality_decision(sw_p, skewness, kurtosis),
        ])
        var_names.append((name, skewness, kurtosis, sw_stat, sw_p))

    if rows:
        _add_apa_table(
            doc,
            ["Değişken", "Çarpıklık", "Basıklık", "Shapiro-Wilk W", "p", "Karar"],
            rows,
        )
        for name, skewness, kurtosis, sw_stat, sw_p in var_names:
            narrative = format_normality_narrative(name, skewness, kurtosis, sw_stat, sw_p)
            if narrative:
                _add_paragraph(doc, narrative)
    else:
        _add_paragraph(doc, "Normallik sonuçları mevcut değil. Lütfen önce Betimsel İstatistikler sayfasını ziyaret edin.")


def _section_group_comparisons(doc: Document, inferential_results) -> None:
    if not inferential_results:
        _add_paragraph(doc, "Grup karşılaştırmaları henüz çalıştırılmamıştır. Bu bölümü dışa aktarmadan önce Grup Karşılaştırmaları modülünü çalıştırın.")
        return

    type_labels = {"ttest": "t-testi", "anova": "ANOVA", "nonparametric": "Parametrik Olmayan"}

    for entry in inferential_results:
        iv = entry.get("iv", "")
        dv = entry.get("dv", "")
        test_type = entry.get("type", "analysis")
        type_label = type_labels.get(test_type, test_type)
        _add_section_heading(doc, f"{iv} × {dv} ({type_label})")

        group_df = entry.get("group_desc_df")
        if isinstance(group_df, pd.DataFrame) and not group_df.empty:
            headers, trows = _dataframe_to_table_rows(group_df)
            _add_paragraph(doc, "Grup betimsel istatistikleri")
            _add_apa_table(doc, headers, trows)

        test_df = entry.get("test_df")
        if isinstance(test_df, pd.DataFrame) and not test_df.empty:
            test_df = test_df.rename(columns={
                "Source": "Kaynak",
                "SS": "KT",
                "df": "sd",
                "MS": "KO",
                "Within": "Grup İçi",
                "Between": "Gruplar Arası",
                "Total": "Toplam",
            })
            source_col = next(
                (c for c in ["Kaynak", "Source"] if c in test_df.columns), None
            )
            if source_col:
                test_df[source_col] = test_df[source_col].replace(
                    {
                        "Within": "Grup İçi",
                        "Total": "Toplam",
                    }
                )
            headers, trows = _dataframe_to_table_rows(test_df)
            _add_paragraph(doc, "Test sonuçları")
            _add_apa_table(doc, headers, trows)

        tukey_df = entry.get("tukey_df")
        if isinstance(tukey_df, pd.DataFrame) and not tukey_df.empty:
            headers, trows = _dataframe_to_table_rows(tukey_df)
            _add_paragraph(doc, "Tukey HSD post-hoc karşılaştırmaları")
            _add_apa_table(doc, headers, trows)

        nonparam = entry.get("nonparametric")
        if nonparam and nonparam.get("summary"):
            # Bold lead-in label, then the full Turkish narrative sentence.
            np_para = doc.add_paragraph()
            np_lead = np_para.add_run("Parametrik olmayan alternatif. ")
            np_lead.bold = True
            np_body = np_para.add_run(nonparam["summary"])
            for run in (np_lead, np_body):
                run.font.name = _FONT_NAME
                run.font.size = _FONT_SIZE_BODY

        # Inline narrative computation
        if test_type == "ttest":
            tdata = entry.get("ttest_data", {})
            if tdata:
                try:
                    narrative = format_ttest_narrative(
                        tdata["var_name"], tdata["group_var"], tdata["group_labels"],
                        tdata["n1"], tdata["n2"],
                        tdata["mean1"], tdata["sd1"],
                        tdata["mean2"], tdata["sd2"],
                        tdata["t"], tdata["df"],
                        tdata["p"], tdata["d"], tdata["levene_p"],
                    )
                    _add_paragraph(doc, narrative)
                except Exception:
                    pass
        elif test_type == "anova":
            adata = entry.get("anova_data", {})
            if adata:
                try:
                    narrative = format_anova_narrative(
                        adata["var_name"], adata["group_var"], adata["k_groups"],
                        adata["F"], adata["df_between"], adata["df_within"],
                        adata["p"], adata["eta2"],
                    )
                    _add_paragraph(doc, narrative)
                except Exception:
                    pass


def _correlation_cell(r: float, p: float) -> str:
    if pd.isna(r) or pd.isna(p):
        return "—"
    stars = "**" if p < 0.01 else ("*" if p < 0.05 else "")
    return f"{format_stat(r, 2)}{stars}"


def _section_correlation(doc: Document, corr_results: dict) -> None:
    method = corr_results.get("method", "Pearson")
    n = corr_results.get("n", "")
    _add_paragraph(doc, f"Korelasyon yöntemi: {method} (N = {n}). * p < .05, ** p < .01.")

    r_df = corr_results.get("matrix")
    p_df = corr_results.get("pmatrix")
    if not isinstance(r_df, pd.DataFrame) or not isinstance(p_df, pd.DataFrame):
        _add_paragraph(doc, "Korelasyon matrisi verisi mevcut değil.")
        return

    cols = list(r_df.columns)
    headers = ["Değişken"] + [_display_name(c) for c in cols]
    rows = []
    for row_name in r_df.index:
        row_cells = [_display_name(row_name)]
        for col_name in r_df.columns:
            row_cells.append(
                _correlation_cell(r_df.loc[row_name, col_name], p_df.loc[row_name, col_name])
            )
        rows.append(row_cells)
    _add_apa_table(doc, headers, rows)

    # Embed heatmap figure
    fig_bytes = st.session_state.get(KEY_CORRELATION_FIGURE)
    if fig_bytes:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(fig_bytes), width=Inches(5.5))
        cap = doc.add_paragraph("Şekil 1. Korelasyon matrisi ısı haritası.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.italic = True
            run.font.size = Pt(10)

    # Inline correlation narratives
    matrix = corr_results.get("matrix")
    pmatrix = corr_results.get("pmatrix")
    if isinstance(matrix, pd.DataFrame) and isinstance(pmatrix, pd.DataFrame):
        var_list = list(matrix.columns)
        for i in range(len(var_list)):
            for j in range(i + 1, len(var_list)):
                v1, v2 = var_list[i], var_list[j]
                r_val = matrix.loc[v1, v2]
                p_val = pmatrix.loc[v1, v2]
                if not (pd.isna(r_val) or pd.isna(p_val)):
                    narrative = format_correlation_narrative(
                        _display_name(v1), _display_name(v2),
                        float(r_val), float(p_val), int(n)
                    )
                    _add_paragraph(doc, narrative)


def _section_moderation(doc: Document, mod_results: dict) -> None:
    x_name = mod_results.get("x", "X")
    y_name = mod_results.get("y", "Y")
    w_name = mod_results.get("w", "W")

    _add_paragraph(doc,
        f"Model özeti (X = {x_name}, Y = {y_name}, W = {w_name}): "
        f"R² = {format_stat(mod_results.get('R2', 0), 3)}, "
        f"adj. R² = {format_stat(mod_results.get('adj_R2', 0), 3)}, "
        f"F({mod_results.get('df1', 0):.0f}, {mod_results.get('df2', 0):.0f}) = "
        f"{mod_results.get('F', 0):.2f}, "
        f"p {format_p(mod_results.get('p_model', 1))}."
    )
    _add_paragraph(doc,
        f"Etkileşim bloğu için ΔR²: {x_name} × {w_name} — "
        f"{format_stat(mod_results.get('delta_R2', 0), 3)}, "
        f"ΔF = {mod_results.get('delta_F', float('nan')):.2f}, "
        f"p {format_p(mod_results.get('delta_p', 1))}."
    )

    ci_ols = mod_results.get("ci_ols", [np.nan, np.nan])
    ci_boot = mod_results.get("ci_boot", [np.nan, np.nan])
    _add_apa_table(doc,
        ["Terim", "B", "SH", "t", "p", "95% GA (OLS)", "95% GA (Bootstrap)"],
        [[
            f"{x_name} × {w_name}",
            format_stat(mod_results.get("b_interaction", 0), 3),
            format_stat(mod_results.get("se_interaction", 0), 3),
            f"{mod_results.get('t_interaction', 0):.2f}",
            format_p(mod_results.get("p_interaction", 1)),
            format_ci(ci_ols[0], ci_ols[1], 3),
            format_ci(ci_boot[0], ci_boot[1], 3),
        ]]
    )

    cond = mod_results.get("conditional_effects", [])
    if cond:
        cond_rows = []
        for row in cond:
            ci = row.get("ci", [np.nan, np.nan])
            cond_rows.append([
                row.get("level", ""),
                format_stat(row.get("b", 0), 3),
                format_stat(row.get("se", 0), 3),
                f"{row.get('t', 0):.2f}",
                format_p(row.get("p", 1)),
                format_ci(ci[0], ci[1], 3),
            ])
        _add_paragraph(doc, "X'in Y üzerindeki koşullu etkileri")
        _add_apa_table(doc, ["W düzeyi", "B", "SH", "t", "p", "95% GA"], cond_rows)

    # Embed interaction plot
    fig_bytes = st.session_state.get(KEY_MODERATION_FIGURE)
    if fig_bytes:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(fig_bytes), width=Inches(5.5))
        cap = doc.add_paragraph(
            f"Şekil 2. {_display_name(w_name)} moderatörlüğünde "
            f"{_display_name(x_name)} → {_display_name(y_name)} ilişkisine ait etkileşim grafiği."
        )
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.italic = True
            run.font.size = Pt(10)

    vif_df = mod_results.get("vif_df")
    if isinstance(vif_df, pd.DataFrame) and not vif_df.empty:
        _add_paragraph(doc, "Varyans enflasyon faktörleri")
        headers, trows = _dataframe_to_table_rows(vif_df)
        _add_apa_table(doc, headers, trows)

    hyp_df = mod_results.get("hypothesis_df")
    if isinstance(hyp_df, pd.DataFrame) and not hyp_df.empty:
        _add_paragraph(doc, "Hipotez özeti")
        headers, trows = _dataframe_to_table_rows(hyp_df)
        _add_apa_table(doc, headers, trows)

    # Inline moderation narratives
    try:
        mod_narrative = format_moderation_narrative(
            x_name, y_name, w_name,
            mod_results.get("R2", 0), mod_results.get("F", 0),
            mod_results.get("df1", 0), mod_results.get("df2", 0),
            mod_results.get("p_model", 1),
            mod_results.get("b_interaction", 0),
            mod_results.get("se_interaction", 0),
            mod_results.get("t_interaction", 0),
            mod_results.get("p_interaction", 1),
            ci_ols[0], ci_ols[1],
            mod_results.get("delta_R2", 0),
        )
        _add_paragraph(doc, mod_narrative)
    except Exception:
        pass

    if len(cond) >= 3:
        try:
            cond_narrative = format_conditional_effects_narrative(
                x_name, y_name, w_name,
                cond[0].get("b", 0), cond[0].get("p", 1),
                cond[1].get("b", 0), cond[1].get("p", 1),
                cond[2].get("b", 0), cond[2].get("p", 1),
            )
            _add_paragraph(doc, cond_narrative)
        except Exception:
            pass


def render():
    st.header("Word'e Aktar")

    final_df = st.session_state.get(KEY_FINAL_DF)
    if final_df is None:
        st.warning("Dışa aktarmadan önce ön işleme adımlarını (Adım 1–4) tamamlayın.")
        return

    composites = st.session_state.get(KEY_COMPOSITE_CONFIG, [])
    col_roles = st.session_state.get(KEY_COL_ROLES, {})
    normality_results = st.session_state.get(KEY_NORMALITY_RESULTS, {})
    inferential_results = st.session_state.get(KEY_INFERENTIAL_RESULTS)
    corr_results = st.session_state.get(KEY_CORRELATION_RESULTS)
    mod_results = st.session_state.get(KEY_MODERATION_RESULTS)

    st.markdown("### Dahil edilecek bölümleri seçin")
    sections = {
        "reliability": st.checkbox("Ölçek Güvenilirliği", value=True, key="export_reliability"),
        "demographics": st.checkbox("Demografik İstatistikler", value=True, key="export_demographics"),
        "descriptives": st.checkbox("Betimsel İstatistikler", value=True, key="export_descriptives"),
        "normality": st.checkbox("Normallik Değerlendirmesi", value=True, key="export_normality"),
        "group_comparisons": st.checkbox(
            f"Grup Karşılaştırmaları {'✅' if inferential_results else '⬜ (henüz çalıştırılmadı)'}",
            value=inferential_results is not None, key="export_group",
        ),
        "correlation": st.checkbox(
            f"Korelasyon Matrisi {'✅' if corr_results else '⬜ (henüz çalıştırılmadı)'}",
            value=corr_results is not None, key="export_correlation",
        ),
        "moderation": st.checkbox(
            f"Moderasyon Analizi {'✅' if mod_results else '⬜ (henüz çalıştırılmadı)'}",
            value=mod_results is not None, key="export_moderation",
        ),
    }

    if st.button("APA Word Belgesi Oluştur", type="primary", key="generate_export_btn"):
        with st.spinner("Belge oluşturuluyor..."):
            doc = _make_document()
            _add_title(doc, "İstatistiksel Analiz Sonuçları")

            if sections["reliability"]:
                _add_section_heading(doc, "Ölçek Güvenilirliği")
                _section_reliability(doc, final_df, composites)

            if sections["demographics"]:
                _add_section_heading(doc, "Demografik İstatistikler")
                _section_demographics(doc, final_df, col_roles)

            if sections["descriptives"]:
                _add_section_heading(doc, "Betimsel İstatistikler")
                _section_descriptives(doc, final_df, composites)

            if sections["normality"]:
                _add_section_heading(doc, "Normallik Değerlendirmesi")
                _section_normality(doc, normality_results, composites, final_df)

            if sections["group_comparisons"]:
                _add_section_heading(doc, "Grup Karşılaştırmaları")
                _section_group_comparisons(doc, inferential_results)

            if sections["correlation"] and corr_results:
                _add_section_heading(doc, "Korelasyon Analizi")
                _section_correlation(doc, corr_results)

            if sections["moderation"] and mod_results:
                _add_section_heading(doc, "Moderasyon Analizi (PROCESS Model 1)")
                _section_moderation(doc, mod_results)

            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = footer_para.add_run(
                f"Analizler Python {platform.python_version()}, "
                f"SciPy {scipy.__version__}, Pingouin {pg.__version__} "
                f"kullanılarak gerçekleştirilmiştir."
            )
            run.font.size = Pt(9)
            run.font.name = _FONT_NAME
            run.italic = True

            # App-version + run-date stamp — traceability of each exported report.
            version_para = doc.add_paragraph()
            version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            version_run = version_para.add_run(
                f"PsychStats v{APP_VERSION} · "
                f"{datetime.date.today().strftime('%d.%m.%Y')}"
            )
            version_run.font.size = Pt(9)
            version_run.font.name = _FONT_NAME
            version_run.italic = True

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)

        st.download_button(
            "⬇ APA Raporunu İndir (.docx)",
            data=buf.getvalue(),
            file_name="psychstats_results.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_docx_btn",
        )
        st.success("Belge hazır. İndirmek için yukarıya tıklayın.")
